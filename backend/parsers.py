import io
from typing import Optional, Tuple
from fastapi import HTTPException, UploadFile
import logging

# Document parsing libraries
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

try:
    import docx
except ImportError:
    docx = None

try:
    from pylatexenc.latex2text import LatexNodes2Text
except ImportError:
    LatexNodes2Text = None

logger = logging.getLogger(__name__)

# Maximum file size: 4MB
MAX_FILE_SIZE = 4_000_000

async def parse_resume(file: Optional[UploadFile], resume_text: Optional[str]) -> Tuple[str, Optional[str]]:
    """
    Parse resume from file or text input.
    
    Args:
        file: Uploaded file (PDF/DOCX/LaTeX)
        resume_text: Direct text input
    
    Returns:
        Tuple of (extracted_text, latex_raw_content)
    """
    # If direct text provided, use it
    if resume_text and resume_text.strip():
        return resume_text.strip(), None
    
    # If no file provided, return empty
    if not file:
        return "", None
    
    try:
        # Read file data
        logger.info(f"Processing file: {file.filename}")
        data = await file.read()
        
        # Check file size
        if len(data) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE // 1_000_000}MB"
            )
        
        # Determine file type
        filename = file.filename.lower() if file.filename else ""
        
        if filename.endswith('.pdf'):
            return await _parse_pdf(data), None
        elif filename.endswith('.docx'):
            return await _parse_docx(data), None
        elif filename.endswith('.tex'):
            return await _parse_latex(data)
        else:
            # Try to guess content type from MIME
            content_type = getattr(file, 'content_type', '') or ''
            if 'pdf' in content_type:
                return await _parse_pdf(data), None
            elif 'word' in content_type or 'document' in content_type:
                return await _parse_docx(data), None
            else:
                raise HTTPException(
                    status_code=415,
                    detail="Unsupported file type. Please upload PDF, DOCX, or LaTeX files."
                )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error parsing file {file.filename}: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Could not parse file. Please ensure it's a valid PDF, DOCX, or LaTeX document."
        )

async def _parse_pdf(data: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        reader = PdfReader(io.BytesIO(data))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
            except Exception as e:
                logger.warning(f"Could not extract text from PDF page {page_num}: {e}")
                continue
        
        if not text_parts:
            raise ValueError("No text content found in PDF")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise ValueError("Could not parse PDF file. The file may be corrupted or password-protected.")

async def _parse_docx(data: bytes) -> str:
    """Extract text from DOCX bytes."""
    if docx is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support not available. Please install python-docx."
        )
    
    try:
        document = docx.Document(io.BytesIO(data))
        text_parts = []
        
        # Extract paragraph text
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract table text
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        if not text_parts:
            raise ValueError("No text content found in DOCX")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise ValueError("Could not parse DOCX file. The file may be corrupted.")

async def _parse_latex(data: bytes) -> Tuple[str, str]:
    """Extract text from LaTeX bytes and return both plain text and raw LaTeX."""
    try:
        # Decode LaTeX source
        latex_raw = data.decode('utf-8', errors='ignore')
        
        if not latex_raw.strip():
            raise ValueError("Empty LaTeX file")
        
        # Convert LaTeX to plain text
        if LatexNodes2Text is None:
            logger.warning("pylatexenc not available, using raw LaTeX")
            # Basic LaTeX cleanup if library not available
            import re
            text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', latex_raw)  # Remove simple commands
            text = re.sub(r'\\[a-zA-Z]+', '', text)  # Remove standalone commands
            text = re.sub(r'[{}]', '', text)  # Remove braces
            text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)  # Remove comments
        else:
            converter = LatexNodes2Text()
            text = converter.latex_to_text(latex_raw)
        
        if not text.strip():
            raise ValueError("Could not extract meaningful text from LaTeX")
        
        return text.strip(), latex_raw
    
    except UnicodeDecodeError:
        raise ValueError("Could not decode LaTeX file. Please ensure it's UTF-8 encoded.")
    except Exception as e:
        logger.error(f"LaTeX parsing error: {e}")
        raise ValueError("Could not parse LaTeX file. Please check the syntax.")

def _clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    
    # Replace multiple whitespace with single space
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # Remove excessive line breaks
    text = re.sub(r'\n+', '\n', text)
    
    return text.strip()